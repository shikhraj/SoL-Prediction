from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
TABLES = OUT / "tables"
PLOTS = OUT / "plots"
REPORT_DIR = OUT / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = REPORT_DIR / "Sleepability_15min_Project_Report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill="F2F4F7") -> None:
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 58, 95)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sleepability MLPR Project Report")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_figure(doc: Document, path: Path, caption: str, width=6.2) -> None:
    if not path.exists():
        p = doc.add_paragraph(f"[Missing figure: {path.name}]")
        p.runs[0].font.color.rgb = RGBColor(155, 28, 28)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def fmt(x, digits=3):
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def add_dataframe_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str] | None = None, max_rows: int | None = None) -> None:
    show = df[columns].copy()
    if max_rows:
        show = show.head(max_rows)
    headers = headers or columns
    table = doc.add_table(rows=1, cols=len(columns))
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                cells[j].text = fmt(val)
            else:
                cells[j].text = str(val)
    style_table(table)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def main() -> None:
    threshold = pd.read_csv(TABLES / "threshold_distribution_comparison.csv")
    ablation = pd.read_csv(TABLES / "feature_group_ablation_results.csv")
    standalone = pd.read_csv(TABLES / "standalone_modality_ablation_results.csv")
    comparison = pd.read_csv(TABLES / "final_model_comparison_15min_success.csv")
    recommendation = pd.read_csv(TABLES / "final_recommendation_summary.csv")
    summary = pd.read_csv(TABLES / "final_execution_summary.csv").iloc[0]
    rf_importance = pd.read_csv(TABLES / "random_forest_feature_importance.csv")
    lr_importance = pd.read_csv(TABLES / "logistic_regression_coefficients.csv")

    doc = Document()
    set_doc_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Personalized Sleepability Prediction from Wearable Data")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Estimating the Probability of Falling Asleep Quickly").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("MLPR Project Report | Fresh 15-minute success analysis").font.color.rgb = RGBColor(85, 85, 85)

    add_callout(
        doc,
        "Executive takeaway",
        "The final model is framed as a sleepability score, P(SOL < 15 minutes). A fresh raw-data rerun tested feature groups through ablation before selecting temporal + previous sleep history. Random Forest was recommended for the final presentation model, while Logistic Regression remains the most interpretable baseline.",
    )

    doc.add_heading("1. Research Goal And Modeling Framing", level=1)
    doc.add_paragraph(
        "The project asks whether information available before bedtime can estimate the probability that a person will fall asleep quickly. The product-facing output is a probability: P(SOL < 15 minutes). This is interpreted as a sleepability score rather than as a retrospective sleep-quality score."
    )
    doc.add_paragraph(
        "The project originally explored delayed sleep onset as a failure event. The final framing uses a success target because it is clearer for a sleepability application: higher probability means better expected ability to fall asleep quickly. Slow onset remains important, so every model is also evaluated from the failure/minority-class perspective."
    )

    doc.add_heading("2. Data Source And Raw Loading", level=1)
    doc.add_paragraph(
        "The fresh pipeline reads raw participant folders from Raw DATA/ifh_affect and detects participant folders automatically. Oura sleep.csv is used as the anchor table because it contains onset_latency and bedtime_start_timestamp. The fresh run detected 24 participant folders, with valid Oura sleep data for 23 participants and 3,997 valid sleep episodes."
    )
    add_bullets(doc, [
        "Raw modalities checked: Oura sleep/activity/readiness/activity level/heart rate/hypnogram, Samsung HRV/pedometer/IMU/PPG/pressure, EMA daily/weekly, Personicle, and assessment files.",
        "Missing files were logged and handled gracefully.",
        "All outputs in this report come from sleepability_15min_full_fresh_analysis, not from the earlier refined pipeline.",
    ])

    doc.add_heading("3. Target Choice: Why 15 Minutes", level=1)
    doc.add_paragraph(
        "Threshold selection is both a statistical and behavioral decision. A 20-minute delayed-SOL threshold is stricter but produces severe class imbalance. A 10-minute threshold is easier to learn but may be too mild and common. A 15-minute threshold provides a middle ground: still behaviorally meaningful while producing enough slow-onset cases for participant-held-out learning."
    )
    add_dataframe_table(
        doc,
        threshold,
        ["threshold_minutes", "slow_onset_count", "quick_onset_count", "slow_onset_rate", "interpretation"],
        ["Threshold", "Slow onset", "Quick onset", "Slow rate", "Interpretation"],
    )
    add_figure(doc, PLOTS / "class_distribution_10_15_20.png", "Figure 1. Slow-onset prevalence under 10-, 15-, and 20-minute thresholds.")

    doc.add_heading("4. Leakage Controls", level=1)
    doc.add_paragraph(
        "The analysis uses only information available at or before bedtime. Same-night sleep outcomes are never used as model inputs for the same target night. This matters because Oura sleep.csv contains many tempting columns that are measured after the night has begun or ended."
    )
    add_bullets(doc, [
        "Forbidden: same-night onset latency, sleep score, sleep duration, sleep stages, efficiency, awake/light/deep/REM totals, midpoint, bedtime_end, sleep HR/HRV, raw timestamps, raw dates, and participant ID as a feature.",
        "Allowed: derived bedtime time features, previous/rolling sleep history shifted within participant, lagged daily summaries, pre-bedtime sensor windows, and EMA records submitted before bedtime.",
        "A feature audit was run before modeling. It passed with zero forbidden model inputs.",
    ])

    doc.add_heading("5. Feature Engineering", level=1)
    doc.add_paragraph(
        "The fresh run engineered multiple leakage-safe feature groups rather than assuming the earlier feature set was best. The groups were designed around ML concepts: start with a simple baseline, add history, then test whether additional modalities provide incremental or standalone signal."
    )
    add_bullets(doc, [
        "Temporal features: bedtime hour, sine/cosine bedtime encoding, day of week, weekend flag, and days since first valid record.",
        "Previous sleep history: previous SOL, previous duration/efficiency/bedtime, rolling SOL/duration/efficiency summaries, bedtime variability, and days since last sleep record.",
        "Oura daily summaries: activity/readiness features lagged to avoid same-day full-day leakage.",
        "Samsung HRV and pedometer: timestamped records restricted to 1h, 3h, and 6h windows before bedtime.",
        "EMA mood: latest daily/weekly EMA submitted before bedtime, with positive affect, negative affect, and anxiety/arousal summaries.",
        "Personicle: optional pre-bedtime context features with interpretable activity summaries.",
    ])

    doc.add_heading("6. Feature Group Ablation: Incremental Value", level=1)
    doc.add_paragraph(
        "The main feature-selection analysis was incremental ablation. Logistic Regression was used as the ablation model because it is stable and interpretable. Each experiment used participant-aware validation and train-fold-only oversampling. The question was: does a modality improve the temporal + previous sleep-history baseline?"
    )
    add_dataframe_table(
        doc,
        ablation,
        ["feature_set", "n_features", "missingness", "pr_auc_success", "pr_auc_slow_onset", "brier_score_success", "f1_slow_onset", "selection_score"],
        ["Feature set", "n", "Missing", "PR success", "PR slow", "Brier", "F1 slow", "Selection"],
    )
    add_figure(doc, PLOTS / "feature_group_ablation_pr_auc_slow_onset.png", "Figure 2. Incremental ablation: slow-onset PR-AUC by feature set.")
    add_figure(doc, PLOTS / "feature_group_ablation_brier.png", "Figure 3. Incremental ablation: Brier score by feature set.")

    doc.add_heading("7. Supplementary Standalone Modality Ablation", level=1)
    doc.add_paragraph(
        "The standalone modality ablation answers a different question: does each modality have predictive signal by itself? This is not the same as incremental ablation. A modality can have standalone signal but still fail to improve a stronger baseline if the signal is redundant, noisy, or too missing."
    )
    add_dataframe_table(
        doc,
        standalone,
        ["feature_set", "n_features", "missingness", "pr_auc_success", "pr_auc_slow_onset", "brier_score_success", "f1_slow_onset", "balanced_accuracy"],
        ["Feature set", "n", "Missing", "PR success", "PR slow", "Brier", "F1 slow", "Bal acc"],
    )
    add_figure(doc, PLOTS / "standalone_modality_ablation_pr_auc_slow_onset.png", "Figure 4. Standalone modality ablation: slow-onset PR-AUC.")
    add_figure(doc, PLOTS / "standalone_vs_incremental_feature_summary.png", "Figure 5. Standalone versus incremental feature summary.")

    doc.add_heading("8. Final Feature Set Selection", level=1)
    doc.add_paragraph(
        "The selected feature set was temporal + previous sleep history. This was not predetermined; it was selected because it had the strongest overall trade-off across participant-held-out performance, slow-onset PR-AUC, calibration, missingness, leakage safety, and interpretability."
    )
    add_bullets(doc, [
        "Full multimodal features did not improve generalization enough to justify the added missingness and complexity.",
        "HRV and Oura showed some signal in certain metrics, but not enough to replace the cleaner temporal/history feature set.",
        "EMA and full multimodal features weakened the main performance story in this dataset.",
        "The final choice supports a simple and defensible MLPR narrative: recent sleep behavior and timing are the strongest predictors of near-term sleepability.",
    ])

    doc.add_heading("9. Validation And Class Imbalance Strategy", level=1)
    doc.add_paragraph(
        "The primary validation method was StratifiedGroupKFold by participant. This preserves participant separation while trying to keep class balance stable across folds. This is stricter than a random row split because the model must generalize to held-out participants."
    )
    doc.add_paragraph(
        "The 15-minute success target is still imbalanced: success is about 73.1% and slow onset is about 26.9%. Oversampling was therefore applied only inside training folds. The held-out test participants keep their natural distribution, preventing test leakage or artificial inflation."
    )

    doc.add_heading("10. Final Model Ladder", level=1)
    doc.add_paragraph(
        "The final ladder tested whether increased model complexity improved the selected feature set: Logistic Regression for interpretability, Random Forest for nonlinear interactions, and CatBoost as a stronger tabular boosting benchmark."
    )
    add_dataframe_table(
        doc,
        comparison,
        ["model_name", "pr_auc_success", "pr_auc_slow_onset", "brier_score_success", "expected_calibration_error_success", "f1_success_tuned", "f1_slow_onset_tuned"],
        ["Model", "PR success", "PR slow", "Brier", "ECE", "F1 success", "F1 slow"],
    )
    add_figure(doc, PLOTS / "model_comparison_pr_auc_success.png", "Figure 6. Final model comparison: PR-AUC for sleep success.")
    add_figure(doc, PLOTS / "model_comparison_pr_auc_slow_onset.png", "Figure 7. Final model comparison: PR-AUC for slow onset.")
    add_figure(doc, PLOTS / "model_comparison_brier.png", "Figure 8. Final model comparison: Brier score for P(success).")
    add_figure(doc, PLOTS / "model_comparison_ece.png", "Figure 9. Final model comparison: expected calibration error.")

    doc.add_heading("11. Model Interpretation", level=1)
    doc.add_paragraph(
        "The interpretation tables indicate which engineered variables the final models relied on most. These importances are descriptive and should not be read as causal effects. Still, they help explain why the final model is behaviorally plausible: it relies heavily on timing regularity and recent sleep history."
    )
    add_dataframe_table(doc, lr_importance.head(8), list(lr_importance.columns), max_rows=8)
    add_figure(doc, PLOTS / "logistic_regression_top_coefficients.png", "Figure 10. Logistic regression top coefficient magnitudes.")
    add_dataframe_table(doc, rf_importance.head(8), list(rf_importance.columns), max_rows=8)
    add_figure(doc, PLOTS / "random_forest_top_features.png", "Figure 11. Random Forest top feature importances.")

    doc.add_heading("12. Final Recommendation", level=1)
    rec_lookup = {row["question"]: row["answer"] for _, row in recommendation.iterrows()}
    doc.add_paragraph(
        f"The recommended final presentation model is {rec_lookup.get('Which model should be used as the final model?', 'Random Forest')}. Random Forest had the best PR-AUC for success, best Brier score, best tuned F1 for success, and best tuned F1 for slow onset. Logistic Regression remains valuable as the most interpretable baseline, and CatBoost achieved the lowest ECE but did not improve discrimination."
    )
    add_bullets(doc, [
        f"Best PR-AUC for success: {summary['best_model_by_PR_AUC_success']}.",
        f"Best PR-AUC for slow onset/failure: {summary['best_model_by_PR_AUC_slow_onset_failure']}.",
        f"Best Brier score: {summary['best_model_by_Brier_score']}.",
        f"Best calibrated model by ECE: {summary['best_calibrated_model']}.",
        f"Recommended final model: {summary['recommended_final_model']}.",
    ])

    doc.add_heading("13. Limitations And Next Steps", level=1)
    doc.add_paragraph(
        "This remains a prototype and is not clinically validated. The dataset is small, observational, and collected from a limited population. Wearable-derived sleep labels are imperfect, and generalization to new users or real deployment would require prospective validation."
    )
    add_bullets(doc, [
        "Add confidence intervals or bootstrap uncertainty around participant-held-out metrics.",
        "Prospectively validate the sleepability score on new participants.",
        "Evaluate calibration by participant and across time.",
        "Investigate why multimodal features add noise or missingness instead of improving generalization.",
        "Consider a user-facing threshold that balances helpful warning with false-alarm cost.",
    ])

    doc.add_heading("14. Reproducibility", level=1)
    doc.add_paragraph("The complete fresh analysis is contained in sleepability_15min_full_fresh_analysis and can be rerun with:")
    p = doc.add_paragraph()
    run = p.add_run("python scripts/run_all.py")
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    doc.add_paragraph(
        "The generated outputs include raw modality availability, feature audit, ablation tables, final model comparison, predictions, plots, logs, and this report."
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
